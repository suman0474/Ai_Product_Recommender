# agentic/deep_agent_document_loader.py
# =============================================================================
# DEEP AGENT DOCUMENT LOADER AND PARSER
# =============================================================================
#
# Loads and parses all standards documents from local storage.
# Extracts sections and analyzes content for memory storage.
#
# =============================================================================

import os
import re
import logging
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime

from docx import Document
from docx.document import Document as DocxDocument

from ..memory import (
    DeepAgentMemory,
    StandardsDocumentAnalysis,
    SectionAnalysis,
    UserContextMemory,
    IdentifiedItemMemory,
    get_relevant_documents_for_product
)

logger = logging.getLogger(__name__)

# =============================================================================
# CONFIGURATION
# =============================================================================

STANDARDS_DIRECTORY = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "chroma_data",
    "Standards"
)

# Section type mapping (normalize section names)
SECTION_TYPE_MAP = {
    "introduction": "introduction",
    "applicable standards": "applicable_standards",
    "applicable standard": "applicable_standards",
    "equipment types": "equipment_types",
    "equipment type": "equipment_types",
    "selection and design criteria": "selection_criteria",
    "selection criteria": "selection_criteria",
    "design criteria": "selection_criteria",
    "accuracy standards and definitions": "accuracy_standards",
    "accuracy standards": "accuracy_standards",
    "installation guidelines": "installation_guidelines",
    "installation guideline": "installation_guidelines",
    "installation best practices": "installation_guidelines",
    "additional installation best practices": "installation_best_practices",
    "calibration and testing": "calibration",
    "calibration": "calibration",
    "testing": "calibration",
    "calibration procedures": "calibration",
    "safety considerations": "safety_considerations",
    "safety consideration": "safety_considerations",
    "safety": "safety_considerations",
    "data sheet requirements": "data_sheet_requirements",
    "data sheet": "data_sheet_requirements",
    "maintenance and inspection": "maintenance",
    "maintenance": "maintenance",
    "maintenance practices": "maintenance",
    "accessories": "accessories",
    "accessory": "accessories",
    "communication technologies": "communication",
    "communication": "communication",
    "signal conditioning": "signal_conditioning",
    "signal conditioning and barriers": "signal_conditioning",
    "sample handling systems": "sample_handling",
    "sample handling": "sample_handling",
    "training and competence": "training",
    "training": "training",
}

# Document type mapping (based on filename)
DOCUMENT_TYPE_MAP = {
    "pressure": "pressure",
    "temperature": "temperature",
    "flow": "flow",
    "level": "level",
    "valves_actuators": "valves",
    "valve": "valves",
    "analytical": "analytical",
    "safety": "safety",
    "control_systems": "control",
    "control": "control",
    "comm_signal": "communication",
    "communication": "communication",
    "calibration_maintenance": "calibration",
    "calibration": "calibration",
    "condition_monitoring": "monitoring",
    "monitoring": "monitoring",
    "accessories": "accessories",
}


# =============================================================================
# DOCUMENT LOADER
# =============================================================================

def load_all_standards_documents(directory: str = None) -> Dict[str, str]:
    """
    Load all standards documents from the local directory.

    Args:
        directory: Path to standards documents directory

    Returns:
        Dictionary mapping filename to full document content
    """
    directory = directory or STANDARDS_DIRECTORY

    if not os.path.exists(directory):
        logger.error(f"Standards directory not found: {directory}")
        return {}

    documents = {}
    files = os.listdir(directory)

    for filename in files:
        if not filename.endswith('.docx'):
            continue

        filepath = os.path.join(directory, filename)

        try:
            doc = Document(filepath)
            content = extract_document_content(doc)
            documents[filename] = content
            logger.info(f"[DocLoader] Loaded: {filename} ({len(content)} chars)")

        except Exception as e:
            logger.error(f"[DocLoader] Failed to load {filename}: {e}")

    logger.info(f"[DocLoader] Loaded {len(documents)} standards documents")
    return documents


def extract_document_content(doc: DocxDocument) -> str:
    """Extract full text content from a Word document."""
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return "\n\n".join(paragraphs)


# =============================================================================
# DOCUMENT PARSER
# =============================================================================

def parse_document_into_sections(doc: DocxDocument) -> List[Dict[str, Any]]:
    """
    Parse a Word document into sections based on heading styles.

    Args:
        doc: python-docx Document object

    Returns:
        List of sections with heading and content
    """
    sections = []
    current_section = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "Normal"

        # Check if this is a heading
        if "Heading" in style_name:
            # Save previous section
            if current_section:
                sections.append(current_section)

            # Start new section
            heading_level = 1
            if "Heading 1" in style_name:
                heading_level = 1
            elif "Heading 2" in style_name:
                heading_level = 2
            elif "Heading 3" in style_name:
                heading_level = 3

            current_section = {
                "heading": text,
                "heading_level": heading_level,
                "style": style_name,
                "content": [],
                "raw_content": ""
            }
        elif current_section:
            current_section["content"].append(text)

    # Save last section
    if current_section:
        sections.append(current_section)

    # Build raw content for each section
    for section in sections:
        section["raw_content"] = "\n".join(section["content"])

    return sections


def get_document_type(filename: str) -> str:
    """Determine document type from filename."""
    filename_lower = filename.lower()

    for key, doc_type in DOCUMENT_TYPE_MAP.items():
        if key in filename_lower:
            return doc_type

    return "general"


def normalize_section_type(section_heading: str) -> str:
    """Normalize section heading to standard section type."""
    heading_lower = section_heading.lower().strip()

    for key, section_type in SECTION_TYPE_MAP.items():
        if key in heading_lower:
            return section_type

    return "other"


# =============================================================================
# CONTENT EXTRACTION
# =============================================================================

def extract_standard_codes(text: str) -> List[str]:
    """
    Extract standard codes from text (e.g., IEC 61508, ISO 1000, API 2540).
    """
    patterns = [
        r'\b(IEC\s*\d+(?:[.-]\d+)*)\b',
        r'\b(ISO\s*\d+(?:[.-]\d+)*)\b',
        r'\b(API\s*\d+(?:[A-Z])?(?:[.-]\d+)*)\b',
        r'\b(ANSI[/\s]+[\w]+\s*\d+(?:[.-]\d+)*)\b',
        r'\b(ISA[/\s-]*\d+(?:[.-]\d+)*)\b',
        r'\b(EN\s*\d+(?:[.-]\d+)*)\b',
        r'\b(NFPA\s*\d+)\b',
        r'\b(ASME\s*[A-Z]*\s*\d+(?:[.-]\d+)*)\b',
        r'\b(IEEE\s*\d+(?:[.-]\d+)*)\b',
        r'\b(NIST\s*[\w]*\s*\d+(?:[.-]\d+)*)\b',
        r'\b(ASTM\s*[A-Z]*\s*\d+(?:[.-]\d+)*)\b',
        r'\b(BS\s*\d+(?:[.-]\d+)*)\b',
        r'\b(DIN\s*\d+(?:[.-]\d+)*)\b',
    ]

    codes = []
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        codes.extend(matches)

    # Clean and deduplicate
    cleaned = list(set([c.strip().upper().replace("  ", " ") for c in codes if c]))
    return sorted(cleaned)


def extract_certifications(text: str) -> List[str]:
    """
    Extract certification codes from text (e.g., SIL2, ATEX, CE).
    """
    patterns = [
        r'\b(SIL\s*[1-4])\b',
        r'\b(ATEX\s*(?:Zone\s*)?[0-2]?(?:/[0-2]+)?)\b',
        r'\b(IECEx)\b',
        r'\b(FM\s*(?:Approved)?)\b',
        r'\b(CSA\s*(?:Certified)?)\b',
        r'\b(UL\s*(?:Listed)?)\b',
        r'\b(CE\s*(?:Mark(?:ed)?)?)\b',
        r'\b(IP[0-9]{2})\b',
        r'\b(NEMA\s*[0-9]+[A-Z]*)\b',
        r'\b(Class\s*[I]+,?\s*Div(?:ision)?\s*[1-2])\b',
        r'\b(TUV\s*(?:Certified)?)\b',
        r'\b(HART\s*(?:\d+)?)\b',
        r'\b(NACE\s*MR\s*\d+)\b',
        r'\b(PED)\b',
    ]

    certifications = []
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        certifications.extend(matches)

    # Clean and deduplicate
    cleaned = list(set([c.strip().upper() for c in certifications if c]))
    return sorted(cleaned)


def extract_equipment_types(text: str) -> List[str]:
    """Extract equipment types mentioned in text."""
    equipment_patterns = [
        r'((?:pressure|differential pressure|absolute pressure)\s+(?:transmitter|gauge|switch|sensor))',
        r'((?:temperature|temp)\s+(?:transmitter|sensor|switch|element))',
        r'((?:thermocouple|RTD|thermistor|pyrometer)s?)',
        r'((?:flow|mass flow|volumetric)\s+(?:meter|transmitter|sensor|switch))',
        r'((?:coriolis|ultrasonic|magnetic|vortex|turbine|orifice)\s+(?:flow\s+)?(?:meter)?)',
        r'((?:level|radar|ultrasonic|guided wave)\s+(?:transmitter|sensor|switch|gauge))',
        r'((?:control|globe|ball|butterfly|gate)\s+valve)',
        r'((?:pneumatic|electric|hydraulic)\s+actuator)',
        r'((?:smart|digital|I/P)\s+positioner)',
        r'((?:pH|conductivity|dissolved oxygen|turbidity)\s+(?:sensor|analyzer|transmitter))',
        r'(gas\s+(?:detector|analyzer|sensor))',
        r'((?:PLC|DCS|RTU|controller)s?)',
        r'((?:vibration|acceleration)\s+(?:sensor|transmitter|monitor))',
    ]

    equipment = []
    for pattern in equipment_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        equipment.extend(matches)

    # Clean and deduplicate
    cleaned = list(set([e.strip().title() for e in equipment if e and len(e) > 3]))
    return sorted(cleaned)


def extract_specifications(text: str) -> Dict[str, str]:
    """
    Extract specification key-value pairs from text.

    Uses tight regex patterns that capture only the technical value portion,
    not surrounding descriptive text. Values are normalized before storage.
    """
    specs = {}

    # Import normalizer for post-processing
    try:
        from ..processing.value_normalizer import ValueNormalizer
        normalizer = ValueNormalizer()
        use_normalizer = True
    except ImportError:
        use_normalizer = False

    # Tight patterns for common specifications - capture ONLY the value
    # Each pattern is designed to stop at natural boundaries
    spec_patterns = {
        # Accuracy: capture percentage with optional tolerance
        "accuracy": r'accuracy[:\s]+([±]?\s*\d+\.?\d*\s*%(?:\s*(?:of\s+(?:span|range|reading|FS))?)?)',
        # Repeatability: similar to accuracy
        "repeatability": r'repeatability[:\s]+([±]?\s*\d+\.?\d*\s*%(?:\s*(?:of\s+(?:span|range|reading|FS))?)?)',
        # Pressure range: capture numeric range with units
        "pressure_range": r'pressure\s+range[:\s]+(-?\d+\.?\d*\s*(?:to|[-–])\s*\d+\.?\d*\s*(?:bar|psi|kPa|MPa|mbar))',
        # Temperature range: capture numeric range with units
        "temperature_range": r'temperature\s+range[:\s]+(-?\d+\.?\d*\s*(?:to|[-–])\s*[+-]?\d+\.?\d*\s*°?[CF])',
        # Flow range: capture numeric range with units
        "flow_range": r'flow\s+range[:\s]+(\d+\.?\d*\s*(?:to|[-–])\s*\d+\.?\d*\s*(?:m³/h|l/min|gpm|kg/h|l/h))',
        # Output signal: capture standard signal types
        "output_signal": r'output\s+(?:signal)?[:\s]+(4-20\s*mA(?:\s*(?:HART|with\s*HART))?|0-10\s*V(?:DC)?|HART\s*\d*|(?:PROFIBUS|Modbus|Foundation\s*Fieldbus)(?:\s+[A-Z]+)?)',
        # Power supply: capture voltage range
        "power_supply": r'power\s+(?:supply)?[:\s]+(\d+\.?\d*\s*(?:to|[-–])\s*\d+\.?\d*\s*V(?:\s*DC)?|\d+\.?\d*\s*V(?:\s*DC)?(?:\s*(?:AC|DC))?)',
        # Response time: capture time value
        "response_time": r'response\s+time[:\s]+([<>]?\s*\d+\.?\d*\s*(?:ms|s|sec|seconds?))',
        # Ingress protection: capture IP rating
        "ingress_protection": r'(?:IP|ingress\s+protection)[:\s]*(IP\s*\d{2}(?:K)?)',
        # Process connection: capture connection type (limited)
        "process_connection": r'process\s+connection[:\s]+(\d+[/"]\d*\s*(?:NPT|BSPT|BSPP|flange|G\d+[/"]\d*)|DN\s*\d+|[A-Z]{2,4}\s*\d+[/"]\d*)',
        # Wetted materials: capture material type (limited to short values)
        "wetted_materials": r'wetted\s+(?:parts|materials)[:\s]+((?:316L?|304)\s*(?:SS|stainless)?|Hastelloy\s*[A-Z]?-?\d*|Inconel\s*\d*|titanium|PTFE)',
        # Ambient temperature: capture range
        "ambient_temperature": r'ambient\s+(?:temperature)?[:\s]+(-?\d+\.?\d*\s*(?:to|[-–])\s*[+-]?\d+\.?\d*\s*°?[CF])',
        # Storage temperature: capture range
        "storage_temperature": r'storage\s+(?:temperature)?[:\s]+(-?\d+\.?\d*\s*(?:to|[-–])\s*[+-]?\d+\.?\d*\s*°?[CF])',
        # Calibration interval: capture time period
        "calibration_interval": r'calibration\s+(?:interval)?[:\s]+(\d+\s*(?:months?|years?|days?))',
        # Sensor type: capture RTD/thermocouple types
        "sensor_type": r'sensor\s+type[:\s]+((?:Pt\s*100|Pt\s*1000|RTD|Type\s*[JKTSERNB]|thermocouple))',
        # Class/grade: capture classification
        "accuracy_class": r'(?:accuracy\s+)?class[:\s]+(Class\s*[A-Z0-9]+|Grade\s*[A-Z0-9]+)',
    }

    for spec_name, pattern in spec_patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            if value and len(value) > 1:
                # Apply normalization if available
                if use_normalizer:
                    result = normalizer.extract_and_validate(value)
                    if result["is_valid"]:
                        specs[spec_name] = result["normalized_value"]
                    else:
                        # Value failed validation - still check if it's short enough
                        if len(value) < 50:
                            specs[spec_name] = value
                else:
                    # No normalizer - basic length check
                    if len(value) < 100:
                        specs[spec_name] = value

    return specs


def extract_guidelines(text: str) -> List[str]:
    """Extract installation/usage guidelines from text."""
    guidelines = []

    # Split into sentences
    sentences = re.split(r'[.!?]\s+', text)

    # Keywords that indicate guidelines
    guideline_keywords = [
        'should', 'shall', 'must', 'recommend', 'require', 'ensure',
        'install', 'mount', 'locate', 'position', 'avoid', 'use',
        'minimum', 'maximum', 'at least', 'no more than'
    ]

    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 20 or len(sentence) > 300:
            continue

        sentence_lower = sentence.lower()
        for keyword in guideline_keywords:
            if keyword in sentence_lower:
                # Clean up the sentence
                clean_sentence = sentence.replace('\n', ' ').strip()
                if clean_sentence and clean_sentence not in guidelines:
                    guidelines.append(clean_sentence)
                break

    return guidelines[:20]  # Limit to 20 guidelines


def extract_key_points(text: str, max_points: int = 10) -> List[str]:
    """Extract key points from section content."""
    key_points = []

    # Look for bullet points or numbered items
    bullet_patterns = [
        r'[•\-\*]\s*([^\n•\-\*]+)',
        r'\d+[.)]\s*([^\n]+)',
    ]

    for pattern in bullet_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            clean_point = match.strip()
            if len(clean_point) > 10 and len(clean_point) < 200:
                key_points.append(clean_point)

    # If no bullet points found, extract first sentences
    if not key_points:
        sentences = re.split(r'[.!?]\s+', text)
        for sentence in sentences[:max_points]:
            sentence = sentence.strip()
            if len(sentence) > 20 and len(sentence) < 200:
                key_points.append(sentence)

    return key_points[:max_points]


# =============================================================================
# DOCUMENT ANALYSIS
# =============================================================================

def analyze_standards_document(
    filename: str,
    doc: DocxDocument
) -> StandardsDocumentAnalysis:
    """
    Analyze a single standards document and extract all relevant information.

    Args:
        filename: Name of the document file
        doc: python-docx Document object

    Returns:
        StandardsDocumentAnalysis with all extracted information
    """
    logger.info(f"[DocAnalysis] Analyzing: {filename}")

    # Parse into sections
    raw_sections = parse_document_into_sections(doc)

    # Get document type
    document_type = get_document_type(filename)

    # Extract full content
    full_content = extract_document_content(doc)

    # Analyze each section
    sections: Dict[str, SectionAnalysis] = {}
    all_standard_codes = []
    all_certifications = []
    all_equipment_types = []
    all_specifications = {}
    all_guidelines = []

    for raw_section in raw_sections:
        section_name = raw_section["heading"]
        section_type = normalize_section_type(section_name)
        raw_content = raw_section["raw_content"]

        # Skip empty sections
        if not raw_content or len(raw_content) < 20:
            continue

        # Extract data from section
        section_codes = extract_standard_codes(raw_content)
        section_certs = extract_certifications(raw_content)
        section_equipment = extract_equipment_types(raw_content)
        section_specs = extract_specifications(raw_content)
        section_guidelines = extract_guidelines(raw_content)
        section_key_points = extract_key_points(raw_content)

        # Build section analysis
        section_analysis: SectionAnalysis = {
            "section_name": section_name,
            "section_type": section_type,
            "raw_content": raw_content,
            "extracted_data": {
                "standard_codes": section_codes,
                "certifications": section_certs,
                "equipment_types": section_equipment,
                "specifications": section_specs,
            },
            "extracted_codes": section_codes,
            "extracted_specs": section_specs,
            "key_points": section_key_points,
            "summary": raw_content[:200] + "..." if len(raw_content) > 200 else raw_content
        }

        sections[section_type] = section_analysis

        # Aggregate data
        all_standard_codes.extend(section_codes)
        all_certifications.extend(section_certs)
        all_equipment_types.extend(section_equipment)
        all_specifications.update(section_specs)
        all_guidelines.extend(section_guidelines)

    # Build document analysis
    analysis: StandardsDocumentAnalysis = {
        "filename": filename,
        "document_type": document_type,
        "full_content": full_content,
        "sections": sections,
        "all_standard_codes": list(set(all_standard_codes)),
        "all_certifications": list(set(all_certifications)),
        "all_equipment_types": list(set(all_equipment_types)),
        "all_specifications": all_specifications,
        "all_guidelines": all_guidelines[:15],  # Limit guidelines
        "total_sections": len(sections),
        "analysis_timestamp": datetime.now().isoformat()
    }

    logger.info(
        f"[DocAnalysis] Completed {filename}: "
        f"{len(sections)} sections, "
        f"{len(all_standard_codes)} standards, "
        f"{len(all_certifications)} certs"
    )

    return analysis


def analyze_all_documents(
    directory: str = None
) -> Dict[str, StandardsDocumentAnalysis]:
    """
    Load and analyze all standards documents.

    Args:
        directory: Path to standards documents directory

    Returns:
        Dictionary mapping filename to document analysis
    """
    directory = directory or STANDARDS_DIRECTORY

    if not os.path.exists(directory):
        logger.error(f"Standards directory not found: {directory}")
        return {}

    analyses = {}
    files = os.listdir(directory)

    for filename in sorted(files):
        if not filename.endswith('.docx'):
            continue

        filepath = os.path.join(directory, filename)

        try:
            doc = Document(filepath)
            analysis = analyze_standards_document(filename, doc)
            analyses[filename] = analysis

        except Exception as e:
            logger.error(f"[DocAnalysis] Failed to analyze {filename}: {e}")

    logger.info(f"[DocAnalysis] Analyzed {len(analyses)} documents")
    return analyses


# =============================================================================
# USER INPUT ANALYSIS
# =============================================================================

def analyze_user_input(user_input: str) -> UserContextMemory:
    """
    Analyze user input to extract domain, process type, and requirements.

    Args:
        user_input: Raw user input string

    Returns:
        UserContextMemory with extracted context
    """
    user_lower = user_input.lower()

    # Domain detection
    domain_keywords = {
        "oil & gas": ["oil", "gas", "refinery", "petroleum", "crude", "lng", "offshore", "pipeline", "hydrocarbon"],
        "chemical": ["chemical", "petrochemical", "reactor", "catalyst", "polymer"],
        "pharmaceutical": ["pharma", "pharmaceutical", "drug", "gmp", "cleanroom", "sterile"],
        "power": ["power", "boiler", "turbine", "generator", "steam", "electricity"],
        "water": ["water", "wastewater", "treatment", "desalination", "potable"],
        "food & beverage": ["food", "beverage", "dairy", "brewery", "sanitary"],
        "pulp & paper": ["pulp", "paper", "mill"],
        "mining": ["mining", "mineral", "ore", "extraction"],
        "metals": ["steel", "aluminum", "metal", "smelting", "foundry"],
    }

    domain = "Industrial"
    for d, keywords in domain_keywords.items():
        for keyword in keywords:
            if keyword in user_lower:
                domain = d.title()
                break

    # Process type detection
    process_keywords = {
        "distillation": ["distillation", "distill", "fractionation", "column"],
        "refining": ["refin", "cracking", "hydrotreating"],
        "reaction": ["reactor", "reaction", "catalyst"],
        "heat exchange": ["heat exchanger", "heating", "cooling", "condenser"],
        "separation": ["separation", "separator", "filter"],
        "compression": ["compressor", "compression"],
        "pumping": ["pump", "pumping"],
        "storage": ["tank", "storage", "vessel"],
        "combustion": ["burner", "furnace", "combustion", "fired"],
    }

    process_type = "General Process"
    for p, keywords in process_keywords.items():
        for keyword in keywords:
            if keyword in user_lower:
                process_type = p.title()
                break

    # Safety requirements detection
    safety_requirements = {}

    # SIL level
    sil_match = re.search(r'sil\s*([1-4])', user_lower)
    if sil_match:
        safety_requirements["sil_level"] = f"SIL {sil_match.group(1)}"
    elif "safety" in user_lower or "critical" in user_lower:
        safety_requirements["sil_level"] = "SIL 2"  # Default for safety-critical

    # Hazardous area
    if any(kw in user_lower for kw in ["hazardous", "explosive", "atex", "zone 0", "zone 1", "zone 2"]):
        safety_requirements["hazardous_area"] = True

        # ATEX zone
        zone_match = re.search(r'zone\s*([0-2])', user_lower)
        if zone_match:
            safety_requirements["atex_zone"] = f"Zone {zone_match.group(1)}"

    # Environmental conditions
    environmental = {}

    if "outdoor" in user_lower:
        environmental["location"] = "outdoor"
    elif "indoor" in user_lower:
        environmental["location"] = "indoor"

    if "corrosive" in user_lower:
        environmental["corrosive"] = True

    if "high temperature" in user_lower or "high temp" in user_lower:
        environmental["high_temperature"] = True

    # ==========================================================
    # CONSTRAINT KEYWORDS - Drive additional standards lookup
    # These trigger searching standards for indirect requirements
    # ==========================================================
    constraint_keywords = []
    
    # Vibration/mechanical constraints
    vibration_keywords = ["vibration", "high vibration", "pump", "compressor", "rotating", "mechanical stress"]
    for kw in vibration_keywords:
        if kw in user_lower:
            constraint_keywords.append("vibration")
            break
    
    # Chemical service constraints (drive material compatibility lookup)
    chemical_services = {
        "hydrogen": ["hydrogen", "h2 service", "hydrogen service"],
        "sour_gas": ["sour gas", "h2s", "hydrogen sulfide", "sour service"],
        "chlorine": ["chlorine", "cl2", "chloride"],
        "ammonia": ["ammonia", "nh3"],
        "oxygen": ["oxygen service", "o2 service", "lox"],
        "caustic": ["caustic", "naoh", "sodium hydroxide"],
        "acid": ["acid", "sulfuric", "hydrochloric", "nitric", "hcl", "h2so4"],
        "seawater": ["seawater", "marine", "offshore", "subsea"],
        "cryogenic": ["cryogenic", "cryo", "lng", "liquid nitrogen", "liquid oxygen"],
    }
    for constraint, keywords in chemical_services.items():
        for kw in keywords:
            if kw in user_lower:
                constraint_keywords.append(constraint)
                break
    
    # Material requirements
    material_keywords = {
        "stainless": ["stainless", "ss316", "316l", "304"],
        "duplex": ["duplex", "super duplex"],
        "inconel": ["inconel", "alloy 625", "alloy 825"],
        "hastelloy": ["hastelloy", "alloy c-276"],
        "titanium": ["titanium", "ti grade"],
        "gold_plating": ["gold plating", "gold plated"],  # For hydrogen service
    }
    for constraint, keywords in material_keywords.items():
        for kw in keywords:
            if kw in user_lower:
                constraint_keywords.append(constraint)
                break
    
    # Physical conditions
    if "abrasive" in user_lower or "slurry" in user_lower or "erosive" in user_lower:
        constraint_keywords.append("abrasive")
    if "viscous" in user_lower or "high viscosity" in user_lower:
        constraint_keywords.append("viscous")
    if "pulsating" in user_lower or "pulsation" in user_lower:
        constraint_keywords.append("pulsation")
    
    # Deduplicate
    constraint_keywords = list(set(constraint_keywords))
    
    if constraint_keywords:
        logger.info(f"[UserAnalysis] Constraint keywords detected: {constraint_keywords}")

    # Key parameters
    key_parameters = {}

    # Temperature range
    temp_match = re.search(r'(\d+)\s*(?:to|[-–])\s*(\d+)\s*(?:°?[CF]|degrees?)', user_input, re.IGNORECASE)
    if temp_match:
        key_parameters["temperature_range"] = f"{temp_match.group(1)} to {temp_match.group(2)}"

    # Pressure range
    pressure_match = re.search(r'(\d+)\s*(?:to|[-–])\s*(\d+)\s*(?:bar|psi|kpa|mpa)', user_input, re.IGNORECASE)
    if pressure_match:
        key_parameters["pressure_range"] = f"{pressure_match.group(1)} to {pressure_match.group(2)}"

    # Extract keywords
    keywords = []
    important_words = re.findall(r'\b[a-zA-Z]{4,}\b', user_input)
    stop_words = {'that', 'this', 'with', 'from', 'have', 'been', 'will', 'would', 'could', 'should', 'need', 'want'}
    keywords = [w.lower() for w in important_words if w.lower() not in stop_words][:20]

    context: UserContextMemory = {
        "raw_input": user_input,
        "domain": domain,
        "process_type": process_type,
        "industry": domain,
        "safety_requirements": safety_requirements,
        "environmental_conditions": environmental,
        "constraint_keywords": constraint_keywords,  # NEW: Drive standards lookup
        "key_parameters": key_parameters,
        "extracted_keywords": list(set(keywords))
    }

    logger.info(f"[UserAnalysis] Domain: {domain}, Process: {process_type}, Constraints: {constraint_keywords}")

    return context


# =============================================================================
# POPULATE MEMORY
# =============================================================================

def populate_memory_with_documents(
    memory: DeepAgentMemory,
    directory: str = None
) -> DeepAgentMemory:
    """
    Load and analyze all documents, then populate memory.

    Args:
        memory: DeepAgentMemory instance to populate
        directory: Path to standards documents

    Returns:
        Populated memory instance
    """
    directory = directory or STANDARDS_DIRECTORY

    logger.info("[PopulateMemory] Starting document analysis...")

    # Analyze all documents
    analyses = analyze_all_documents(directory)

    # Store in memory
    for filename, analysis in analyses.items():
        memory.store_standards_analysis(filename, analysis)

    # Build section index (cross-document)
    section_types = [
        "applicable_standards",
        "equipment_types",
        "selection_criteria",
        "safety_considerations",
        "installation_guidelines",
        "calibration"
    ]

    for section_type in section_types:
        section_data = {
            "documents": [],
            "all_content": [],
            "all_codes": [],
            "all_specs": {}
        }

        for filename, analysis in analyses.items():
            sections = analysis.get("sections", {})
            if section_type in sections:
                section = sections[section_type]
                section_data["documents"].append(filename)
                section_data["all_content"].append(section.get("raw_content", ""))
                section_data["all_codes"].extend(section.get("extracted_codes", []))
                section_data["all_specs"].update(section.get("extracted_specs", {}))

        # Deduplicate codes
        section_data["all_codes"] = list(set(section_data["all_codes"]))

        memory.store_section_index(section_type, section_data)

    logger.info(f"[PopulateMemory] Completed: {memory.documents_loaded} docs, {memory.sections_analyzed} sections")

    return memory


def populate_memory_with_items(
    memory: DeepAgentMemory,
    instruments: List[Dict[str, Any]],
    accessories: List[Dict[str, Any]],
    user_context: Dict[str, Any] = None
) -> DeepAgentMemory:
    """
    Populate memory with identified instruments and accessories.

    Args:
        memory: DeepAgentMemory instance
        instruments: List of identified instruments
        accessories: List of identified accessories
        user_context: Optional user context with constraint_keywords for standards lookup

    Returns:
        Populated memory instance
    """
    # Extract constraint keywords from user context
    constraint_keywords = []
    if user_context:
        constraint_keywords = user_context.get("constraint_keywords", [])
        if constraint_keywords:
            logger.info(f"[PopulateMemory] Using constraint keywords: {constraint_keywords}")
    
    # Process instruments
    for i, inst in enumerate(instruments):
        # Support multiple key names for product type
        product_type = (
            inst.get("product_type") or
            inst.get("product_name") or
            inst.get("category") or
            "Unknown Instrument"
        )
        # Pass constraint keywords for context-aware document retrieval
        relevant_docs = get_relevant_documents_for_product(product_type, constraint_keywords)

        item: IdentifiedItemMemory = {
            "item_id": f"inst_{i+1}",
            "item_type": "instrument",
            "product_type": product_type,
            "category": inst.get("category", ""),
            "quantity": inst.get("quantity", 1),
            "user_specifications": inst.get("user_specifications") or inst.get("specifications", {}),
            "relevant_documents": relevant_docs,
            "related_instrument": None
        }

        memory.store_identified_item(item)
        memory.store_product_type_mapping(product_type, relevant_docs)

    # Process accessories
    for i, acc in enumerate(accessories):
        # Support multiple key names for product type
        product_type = (
            acc.get("product_type") or
            acc.get("accessory_name") or
            acc.get("category") or
            "Unknown Accessory"
        )
        # Pass constraint keywords for context-aware document retrieval
        relevant_docs = get_relevant_documents_for_product(product_type, constraint_keywords)

        item: IdentifiedItemMemory = {
            "item_id": f"acc_{i+1}",
            "item_type": "accessory",
            "product_type": product_type,
            "category": acc.get("category", ""),
            "quantity": acc.get("quantity", 1),
            "user_specifications": acc.get("user_specifications") or acc.get("specifications", {}),
            "relevant_documents": relevant_docs,
            "related_instrument": acc.get("related_instrument") or acc.get("for_instrument")
        }

        memory.store_identified_item(item)
        memory.store_product_type_mapping(product_type, relevant_docs)

    logger.info(
        f"[PopulateMemory] Added {len(instruments)} instruments, "
        f"{len(accessories)} accessories (constraints: {constraint_keywords})"
    )

    return memory


# =============================================================================
# EXPORTS
# =============================================================================

__all__ = [
    "STANDARDS_DIRECTORY",
    "load_all_standards_documents",
    "parse_document_into_sections",
    "analyze_standards_document",
    "analyze_all_documents",
    "analyze_user_input",
    "populate_memory_with_documents",
    "populate_memory_with_items",
    "extract_standard_codes",
    "extract_certifications",
    "extract_equipment_types",
    "extract_specifications",
    "extract_guidelines",
]
